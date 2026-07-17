"""Convert a literature review markdown note to a formatted Word doc.

Defaults to the neurosymbolic review this was written for; pass an input
path (and optionally an output path) to convert any other note:

    python scripts/md_to_docx.py [input.md] [output.docx]
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO_ROOT = Path(__file__).resolve().parent.parent

DEFAULT_INPUT_MD = (
    REPO_ROOT
    / "Library"
    / "Inference-Time Reasoning Algorithms"
    / "2026-05-15 Literature Review - Neurosymbolic AI and Formal Verification for AWS Workflows.md"
)
DEFAULT_OUTPUT_DOCX = (
    REPO_ROOT
    / "2026-05-15 Neurosymbolic AI & Formal Verification - Literature Review.docx"
)

# ── Colours ──────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x3A, 0x5C)   # H1
SLATE  = RGBColor(0x2C, 0x4A, 0x6E)   # H2
TEAL   = RGBColor(0x1A, 0x6A, 0x72)   # H3
BODY   = RGBColor(0x1A, 0x1A, 0x2E)   # body text
GREY   = RGBColor(0x60, 0x60, 0x70)   # meta / label text
ACCENT = RGBColor(0x1A, 0x6A, 0x72)   # table header bg

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), kwargs.get(edge, "none"))
        tag.set(qn("w:sz"), kwargs.get("sz", "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "CCCCCC"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def strip_md_inline(text: str) -> list[tuple[str, dict]]:
    """
    Parse a markdown inline string and return a list of (text, style) tuples.
    Handles **bold**, *italic*, ***bold-italic***, [text](url) links,
    and combinations like *[text](url)*.
    """
    runs = []
    # Regex that captures: bold+italic, bold, italic, link, plain
    pattern = re.compile(
        r'\*\*\*(.+?)\*\*\*'      # bold+italic
        r'|\*\*(.+?)\*\*'          # bold
        r'|\*\[(.+?)\]\((.+?)\)\*' # italic link  *[text](url)*
        r'|\[(.+?)\]\((.+?)\)'     # plain link [text](url)
        r'|\*(.+?)\*'              # italic
        r'|([^*\[\]]+)'            # plain text
    )
    for m in pattern.finditer(text):
        bi, b, lt, lu, lnk_t, lnk_u, it, plain = m.groups()
        if bi:
            runs.append((bi, {"bold": True, "italic": True}))
        elif b:
            runs.append((b, {"bold": True}))
        elif lt:
            runs.append((lt, {"italic": True, "url": lu}))
        elif lnk_t:
            runs.append((lnk_t, {"url": lnk_u}))
        elif it:
            runs.append((it, {"italic": True}))
        elif plain:
            runs.append((plain, {}))
    return runs


def add_inline_runs(para, text: str, base_size: int = 11,
                    base_color: RGBColor = BODY, base_bold: bool = False):
    """Render inline-markdown text into a paragraph with proper runs."""
    for run_text, style in strip_md_inline(text):
        run = para.add_run(run_text)
        run.font.size = Pt(base_size)
        run.font.color.rgb = base_color
        run.bold = style.get("bold", base_bold)
        run.italic = style.get("italic", False)
        if "url" in style:
            # Add hyperlink
            run.font.color.rgb = RGBColor(0x1A, 0x6A, 0x72)
            run.font.underline = True
            # Word hyperlink via relationship
            rId = para.part.target_ref(style["url"])
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), rId)
            hyperlink.append(run._r)
            para._p.append(hyperlink)
            # remove the run from the paragraph (it's inside the hyperlink now)
            continue
    return para


def add_hyperlink(para, text: str, url: str,
                  color: RGBColor = RGBColor(0x1A, 0x6A, 0x72),
                  size: int = 11, bold: bool = False, italic: bool = False):
    """Add a clickable hyperlink run to a paragraph."""
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    style_el = OxmlElement("w:rStyle")
    style_el.set(qn("w:val"), "Hyperlink")
    rPr.append(style_el)
    # color
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    rPr.append(color_el)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    rPr.append(sz)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    hyperlink.append(run)
    para._p.append(hyperlink)


def add_plain_run(para, text: str, size: int = 11,
                  color: RGBColor = BODY, bold: bool = False, italic: bool = False):
    r = para.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


def render_inline(para, text: str, size: int = 11,
                  base_color: RGBColor = BODY, base_bold: bool = False,
                  base_italic: bool = False):
    """
    Render markdown inline markup into paragraph runs.
    Handles: **bold**, *italic*, [text](url), *[text](url)*, ***bold+italic***
    """
    pattern = re.compile(
        r'\*\*\*(.+?)\*\*\*'
        r'|\*\*(.+?)\*\*'
        r'|\*\[(.+?)\]\((.+?)\)\*'
        r'|\[(.+?)\]\((.+?)\)'
        r'|\*(.+?)\*'
        r'|([^*\[\]]+)'
    )
    for m in pattern.finditer(text):
        bi, b, lt, lu, lnk_t, lnk_u, it, plain = m.groups()
        if bi:
            add_plain_run(para, bi, size=size, color=base_color, bold=True, italic=True)
        elif b:
            add_plain_run(para, b, size=size, color=base_color,
                          bold=True, italic=base_italic)
        elif lt:  # italic link
            add_hyperlink(para, lt, lu, size=size, italic=True)
        elif lnk_t:
            add_hyperlink(para, lnk_t, lnk_u, size=size,
                          bold=base_bold, italic=base_italic)
        elif it:
            add_plain_run(para, it, size=size, color=base_color,
                          bold=base_bold, italic=True)
        elif plain:
            add_plain_run(para, plain, size=size, color=base_color,
                          bold=base_bold, italic=base_italic)


# ── Document setup ────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name  = "Calibri"
    style.font.size  = Pt(11)
    style.font.color.rgb = BODY

    return doc


# ── Paragraph builders ────────────────────────────────────────────────────────

def add_title(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(24)
    r.font.color.rgb = NAVY
    r.bold = True
    return p


def add_meta(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    render_inline(p, text, size=10, base_color=GREY)
    return p


def add_h1(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(15)
    r.font.color.rgb = NAVY
    r.bold = True
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A3A5C")
    pb.append(bottom)
    pPr.append(pb)
    return p


def add_h2(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(13)
    r.font.color.rgb = SLATE
    r.bold = True
    return p


def add_h3(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(11.5)
    r.font.color.rgb = TEAL
    r.bold = True
    return p


def add_body(doc, text: str, indent: float = 0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Inches(indent)
    render_inline(p, text, size=11, base_color=BODY)
    return p


def add_bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    render_inline(p, text, size=11, base_color=BODY)
    return p


def add_numbered(doc, text: str, num: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    add_plain_run(p, f"{num}. ", size=11, color=BODY, bold=True)
    render_inline(p, text, size=11, base_color=BODY)
    return p


def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pb.append(bottom)
    pPr.append(pb)
    return p


# ── Table builder ─────────────────────────────────────────────────────────────

def add_table(doc, header_row: list[str], rows: list[list[str]],
              col_widths: list[float] | None = None):
    n_cols = len(header_row)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)

    # Header row
    hdr = tbl.rows[0]
    for i, text in enumerate(header_row):
        cell = hdr.cells[i]
        set_cell_bg(cell, "1A3A5C")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(text)
        r.font.name  = "Calibri"
        r.font.size  = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.bold = True

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg  = "F4F7FA" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            render_inline(p, cell_text, size=10, base_color=BODY)

    doc.add_paragraph()  # spacing after table
    return tbl


# ── Markdown parser ───────────────────────────────────────────────────────────

def parse_md_table(block: list[str]) -> tuple[list[str], list[list[str]]]:
    """Parse a markdown pipe table into header + rows."""
    header = [c.strip() for c in block[0].strip("|").split("|")]
    rows   = []
    for line in block[2:]:  # skip separator line
        if not line.strip():
            continue
        rows.append([c.strip() for c in line.strip("|").split("|")])
    return header, rows


def render_md(doc: Document, md_text: str):
    lines = md_text.splitlines()
    i = 0
    in_table = False
    table_block: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Flush table
        if in_table and not line.startswith("|"):
            header, rows = parse_md_table(table_block)
            n = len(header)
            # Heuristic column widths
            if n == 3:
                widths = [2.2, 3.0, 1.8]
            elif n == 4:
                widths = [2.6, 1.1, 0.55, 2.2]
            else:
                widths = [7.0 / n] * n
            add_table(doc, header, rows, widths)
            in_table = False
            table_block = []

        # HR
        if line.strip() in ("---", "***", "___"):
            add_rule(doc)
            i += 1
            continue

        # Table
        if line.startswith("|"):
            if not in_table:
                in_table = True
                table_block = []
            table_block.append(line)
            i += 1
            continue

        # H1 (# Title)
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            add_title(doc, text)
            i += 1
            continue

        # H2
        if line.startswith("## "):
            text = line[3:].strip()
            add_h1(doc, text)
            i += 1
            continue

        # H3
        if line.startswith("### "):
            text = line[4:].strip()
            add_h2(doc, text)
            i += 1
            continue

        # H4
        if line.startswith("#### "):
            text = line[5:].strip()
            add_h3(doc, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            add_numbered(doc, m.group(2), int(m.group(1)))
            i += 1
            continue

        # Bullet
        if line.startswith("- "):
            add_bullet(doc, line[2:])
            i += 1
            continue

        # Meta lines (bold key: value at top)
        if line.startswith("**") and ":**" in line:
            add_meta(doc, line)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Body paragraph
        add_body(doc, line)
        i += 1

    # Flush remaining table
    if in_table and table_block:
        header, rows = parse_md_table(table_block)
        n = len(header)
        if n == 3:
            widths = [2.2, 3.0, 1.8]
        elif n == 4:
            widths = [2.6, 1.1, 0.55, 2.2]
        else:
            widths = [7.0 / n] * n
        add_table(doc, header, rows, widths)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    input_md = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_INPUT_MD
    if len(sys.argv) > 2:
        output_docx = Path(sys.argv[2])
    elif len(sys.argv) > 1:
        output_docx = input_md.with_suffix(".docx")
    else:
        output_docx = DEFAULT_OUTPUT_DOCX

    if not input_md.is_file():
        raise SystemExit(f"Input note not found: {input_md}")

    md = input_md.read_text(encoding="utf-8")

    doc = build_doc()
    render_md(doc, md)
    doc.save(output_docx)
    print(f"Saved: {output_docx}")


if __name__ == "__main__":
    main()
